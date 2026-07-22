import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
os.environ.setdefault("URDU_UNICODER_DISABLE_UPDATE_CHECK", "1")
sys.path.insert(0, str(Path(__file__).resolve().parent / "app"))

from PySide6.QtCore import QMarginsF
from PySide6.QtGui import QPageLayout, QPageSize, QTextDocument
from PySide6.QtPrintSupport import QPrinter
from PySide6.QtWidgets import QApplication, QScrollArea, QToolBar

from main import (
    APP_NAME, APP_VERSION, ICON_PATH, LOGO_PATH, ExtractWorker, LayoutSettings,
    MainWindow, extract_docx_text, version_tuple,
)


class UnicodeRecoveryTests(unittest.TestCase):
    def setUp(self):
        self.worker = ExtractWorker("", 1, 1, LayoutSettings())

    def test_presentation_forms_and_visual_order_are_recovered(self):
        source = "ﮐﺎ اس اور ﺻﺒﺎ اور ﺗﮭﮯ رﮨﮯ ﺑﺞ دس ﮐﮯ رات"
        self.assertEqual(self.worker.clean_line(source), "رات کے دس بج رہے تھے اور صبا اور اس کا")

    def test_ellipsis_is_normalized(self):
        source = ". . . ﺗﮭﮯ ﮨﻮﺋﮯ ﺑﯿﭩﮭﮯ ﻣﯿﮟ ﺑﯿﮉروم اﭘﻨﮯ"
        self.assertEqual(self.worker.clean_line(source), "اپنے بیڈروم میں بیٹھے ہوئے تھے…")

    def test_join_lines_keeps_real_paragraph_boundaries(self):
        text = "first line\nsecond line\n\nnext paragraph"
        self.assertEqual(MainWindow._join_wrapped_lines(text), "first line second line\n\nnext paragraph")


class UiAndPdfTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.app = QApplication.instance() or QApplication([])

    def test_professional_settings_panel_is_scrollable(self):
        window = MainWindow()
        scroll = window.findChild(QScrollArea, "settingsScroll")
        self.assertIsNotNone(scroll)
        self.assertTrue(scroll.widgetResizable())
        self.assertEqual(APP_NAME, "Urdu Unicoder")
        window.close()

    def test_logo_and_advanced_editor_are_available(self):
        window = MainWindow()
        self.assertTrue(LOGO_PATH.exists())
        self.assertTrue(ICON_PATH.exists())
        self.assertFalse(window.windowIcon().isNull())
        self.assertIsNotNone(window.findChild(QScrollArea, "settingsScroll"))
        self.assertIsNotNone(window.findChild(QToolBar, "editorToolbar"))
        window.close()

    def test_update_manifest_matches_application_version(self):
        manifest = json.loads((Path(__file__).parent / "version.json").read_text(encoding="utf-8"))
        self.assertEqual(manifest["version"], APP_VERSION)
        self.assertGreater(version_tuple("1.10.0"), version_tuple("1.9.9"))

    def test_unicode_find_replace_and_normalization(self):
        window = MainWindow()
        window.editor.setPlainText("کتاب اچھی ہے۔ کتاب مفید ہے۔")
        window.find_input.setText("کتاب")
        window.replace_input.setText("تحریر")
        window.replace_all()
        self.assertEqual(window.editor.toPlainText(), "تحریر اچھی ہے۔ تحریر مفید ہے۔")
        window.editor.setPlainText("ﮐﺘﺎﺏ")
        window.normalize_editor_unicode()
        self.assertEqual(window.editor.toPlainText(), "کتاب")
        window.editor.clear()
        window.close()

    def test_clipboard_paste_converts_presentation_forms(self):
        window = MainWindow()
        QApplication.clipboard().setText("ﮐﺘﺎﺏ اردو")
        window.paste_and_convert_unicode()
        self.assertEqual(window.editor.toPlainText(), "کتاب اردو")
        window.editor.clear()
        window.close()

    def test_word_import_preserves_paragraph_table_order(self):
        from docx import Document

        path = Path(tempfile.gettempdir()) / "urdu_unicoder_import_test.docx"
        try:
            document = Document()
            document.add_heading("کتاب کا عنوان", level=1)
            document.add_paragraph("پہلا پیراگراف")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "نام"
            table.cell(0, 1).text = "تفصیل"
            document.add_paragraph("آخری پیراگراف")
            document.save(path)
            imported = extract_docx_text(str(path))
            expected = ["کتاب کا عنوان", "پہلا پیراگراف", "نام | تفصیل", "آخری پیراگراف"]
            positions = [imported.index(value) for value in expected]
            self.assertEqual(positions, sorted(positions))
        finally:
            path.unlink(missing_ok=True)

    def test_current_pyside_pdf_margin_api(self):
        output = Path(tempfile.gettempdir()) / "urdu_unicoder_margin_test.pdf"
        try:
            printer = QPrinter(QPrinter.HighResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(str(output))
            printer.setPageSize(QPageSize(QPageSize.A5))
            printer.setPageMargins(QMarginsF(15, 18, 20, 18), QPageLayout.Millimeter)
            document = QTextDocument("Urdu Unicoder PDF export test")
            document.print_(printer)
            self.assertTrue(output.exists())
            self.assertGreater(output.stat().st_size, 0)
        finally:
            output.unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()
